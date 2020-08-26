import docx


def parse(document):
    text = []
    doc = docx.Document(document.file)
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for cell in cells:
                text.append(cell.text)
    return "\n ".join(text)
